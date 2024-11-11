from openai import OpenAI
import pandas as pd
from docx import Document
import os
from datetime import datetime

# Your OpenAI API key - Replace this with your actual API key
OPENAI_API_KEY = "my_api_key"  # Replace this with your actual API key

# Initialize OpenAI client
client = OpenAI(api_key=OPENAI_API_KEY)

# Read patient data from Excel file
data = pd.read_excel('patient_data.xlsx')

def generate_report(patient_info):
    try:
        # Construct the prompt
        prompt = f"""
        You are a psychiatrist generating a consultation report. Use the following patient information to create a structured report:

        Patient ID: {patient_info['Patient_ID']}
        Name: {patient_info['Name']}
        Age: {patient_info['Age']}
        Gender: {patient_info['Gender']}
        Symptoms: {patient_info['Symptoms']}
        Medical History: {patient_info['Medical_History']}
        Diagnosis: {patient_info['Diagnosis']}
        Consultation Notes: {patient_info['Consultation_Notes']}

        The report should include the following sections:
        - Introduction
        - Patient History
        - Mental Status Examination
        - Diagnosis
        - Treatment Plan
        - Recommendations

        Use appropriate psychiatric terminology and maintain a professional tone.
        """

        # Call the OpenAI API
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",  # Changed to gpt-3.5-turbo as it's more cost-effective
            messages=[
                {"role": "system", "content": "You are a psychiatrist."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=1500,
            temperature=0.7
        )

        # Extract the assistant's reply
        report = response.choices[0].message.content.strip()
        return report

    except Exception as e:
        print(f"Error generating report for Patient ID: {patient_info['Patient_ID']}")
        print(e)
        return None

def main():
    # Create an output directory for the reports
    output_dir = 'reports'
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    # Iterate over each patient record
    for index, row in data.iterrows():
        patient_info = row.to_dict()
        report_text = generate_report(patient_info)

        if report_text:
            # Create a Word document
            document = Document()
            document.add_heading('Psychiatric Consultation Report', 0)
            
            # Add timestamp
            timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            document.add_paragraph(f"Generated on: {timestamp}")
            
            document.add_paragraph(report_text)

            # Save the report
            patient_id = patient_info['Patient_ID']
            filename = f'report_{patient_id}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
            document.save(os.path.join(output_dir, filename))

            print(f"Report generated for Patient ID: {patient_id}")
        else:
            print(f"Failed to generate report for Patient ID: {patient_info['Patient_ID']}")

if __name__ == "__main__":
    main()
